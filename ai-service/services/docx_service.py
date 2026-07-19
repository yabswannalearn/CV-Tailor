import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def strip_latex_formatting(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r'\\textbf\{(.*?)\}', r'\1', text)
    text = re.sub(r'\\textit\{(.*?)\}', r'\1', text)
    text = re.sub(r'\\emph\{(.*?)\}', r'\1', text)
    text = re.sub(r'\\href\{.*?\}\{(.*?)\}', r'\1', text)
    text = re.sub(r'\\url\{(.*?)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+\{?\}?', '', text)
    text = text.replace(r'\&', '&').replace(r'\%', '%').replace(r'\$', '$').replace(r'\_', '_').replace(r'\|', '|')
    return text.replace('{', '').replace('}', '').strip()

def convert_latex_to_docx(latex_code: str) -> bytes:
    doc = Document()
    
    # 1. Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    # 2. Base Normal Style Configuration
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    lines = latex_code.split('\n')
    in_document = False
    
    for line in lines:
        line_str = line.strip()
        if r"\begin{document}" in line_str:
            in_document = True
            continue
        if r"\end{document}" in line_str:
            break
        if not in_document or not line_str or line_str.startswith('%'):
            continue
            
        # Parse Section Headers \section{...}
        section_match = re.search(r'\\section\*?\{(.*?)\}', line_str)
        if section_match:
            raw_title = section_match.group(1)
            title = strip_latex_formatting(raw_title)
            h = doc.add_heading(title.upper(), level=2)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(2)
            for run in h.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x22, 0x55, 0x00) # Subtle accent color
            continue

        # Parse Subheadings / Headings \resumeSubheading or \resumeProjectHeading
        if r"\resumeSubheading" in line_str or r"\resumeProjectHeading" in line_str:
            sub_matches = re.findall(r'\{(.*?)\}', line_str)
            if sub_matches:
                clean_parts = [strip_latex_formatting(p) for p in sub_matches if strip_latex_formatting(p)]
                heading_line = " | ".join(clean_parts)
                if heading_line:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(1)
                    run = p.add_run(heading_line)
                    run.bold = True
                    run.font.size = Pt(10.5)
            continue
            
        # Parse Bullet Items \item or \resumeItem
        if line_str.startswith(r'\item') or r'\resumeItem' in line_str:
            item_text = re.sub(r'\\(item|resumeItem)\s*\{?', '', line_str)
            item_text = strip_latex_formatting(item_text)
            if item_text:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1.5)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(item_text)
                run.font.size = Pt(10)
            continue
            
        # Regular Paragraph
        clean_text = strip_latex_formatting(line_str)
        if clean_text and not clean_text.startswith('\\'):
            p = doc.add_paragraph(clean_text)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
